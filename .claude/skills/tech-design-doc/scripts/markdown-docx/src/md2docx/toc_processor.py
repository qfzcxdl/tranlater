#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
import unicodedata
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def slugify(text: str) -> str:
    text = unicodedata.normalize("NFKD", text)
    text = re.sub(r"[^\w\s-]", "", text.lower())
    text = re.sub(r"[-\s]+", "-", text).strip("-")
    return text


def add_bookmark_to_paragraph(paragraph, bookmark_name: str, bookmark_id: int):
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bookmark_id))
    bookmark_start.set(qn("w:name"), bookmark_name)

    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bookmark_id))

    p_element = paragraph._p
    if p_element is not None and len(p_element) > 0:
        p_element.insert(0, bookmark_start)
        p_element.append(bookmark_end)


def create_internal_hyperlink(paragraph, bookmark_name: str, display_text: str):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)

    run = OxmlElement("w:r")

    run_props = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_props.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)

    run.append(run_props)

    text = OxmlElement("w:t")
    text.text = display_text
    run.append(text)

    hyperlink.append(run)

    paragraph._p.append(hyperlink)


def process_docx_toc(docx_path: str, output_path: str = None):
    if output_path is None:
        output_path = docx_path

    doc = Document(docx_path)

    heading_bookmarks = {}
    bookmark_id = 0

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            heading_text = paragraph.text.strip()
            if heading_text:
                bookmark_name = slugify(heading_text)
                if not bookmark_name:
                    bookmark_name = f"heading_{bookmark_id}"

                if bookmark_name in heading_bookmarks:
                    bookmark_name = f"{bookmark_name}_{bookmark_id}"

                add_bookmark_to_paragraph(paragraph, bookmark_name, bookmark_id)
                heading_bookmarks[heading_text] = bookmark_name
                bookmark_id += 1

    toc_link_pattern = re.compile(r"\[([^\]]+)\]\(#([^)]+)\)")

    for paragraph in doc.paragraphs:
        text = paragraph.text
        matches = list(toc_link_pattern.finditer(text))

        if not matches:
            continue

        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading"):
            continue

        level = 0
        if text.strip().startswith("-"):
            leading_spaces = len(text) - len(text.lstrip())
            level = leading_spaces // 2

        paragraph.clear()

        if level > 0:
            paragraph.paragraph_format.left_indent = Pt(18 * level)

        last_end = 0
        for match in matches:
            if match.start() > last_end:
                prefix = text[last_end : match.start()].lstrip("- ")
                if prefix:
                    paragraph.add_run(prefix)

            display_text = match.group(1)
            anchor = match.group(2)

            bookmark_name = None
            for heading_text, bm_name in heading_bookmarks.items():
                if slugify(heading_text) == anchor or bm_name == anchor:
                    bookmark_name = bm_name
                    break

            if bookmark_name is None:
                bookmark_name = anchor

            create_internal_hyperlink(paragraph, bookmark_name, display_text)

            last_end = match.end()

        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    doc.save(output_path)
    return len(heading_bookmarks)


def process_all_docx_in_directory(input_dir: str, output_dir: str = None):
    from pathlib import Path

    input_path = Path(input_dir)
    output_path = Path(output_dir) if output_dir else input_path

    output_path.mkdir(parents=True, exist_ok=True)

    docx_files = list(input_path.glob("*.docx"))
    results = []

    for docx_file in docx_files:
        output_file = output_path / docx_file.name
        try:
            bookmark_count = process_docx_toc(str(docx_file), str(output_file))
            results.append((docx_file.name, bookmark_count, True))
            print(f"✅ {docx_file.name}: {bookmark_count} bookmarks added")
        except Exception as e:
            results.append((docx_file.name, 0, False))
            print(f"❌ {docx_file.name}: {e}")

    return results


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Add clickable TOC to DOCX files")
    parser.add_argument("-i", "--input", default="output", help="Input directory")
    parser.add_argument("-o", "--output", default=None, help="Output directory")

    args = parser.parse_args()

    print()
    print("🔗 Adding clickable TOC links to DOCX files")
    print("=" * 60)
    print()

    results = process_all_docx_in_directory(args.input, args.output or args.input)

    print()
    print("=" * 60)
    success = sum(1 for _, _, ok in results if ok)
    print(f"📊 Processed: {success}/{len(results)} files")
